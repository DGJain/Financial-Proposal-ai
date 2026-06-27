"""Build the BITS WILP Mid-Semester Report (.docx) for the dissertation.

Generates a Word document mirroring the institute's mid-semester report layout
for Garvit D Jain (2024AA05088) — "Agentic SLM-Based Multi-Source Intelligent
Proposal Generation System for Enterprise Consulting Using Adaptive Model
Orchestration". Content is grounded in the working platform in this repo.

Run:  backend\.venv\Scripts\python.exe infra\scripts\dev\build_midsem_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # deep blue, section headings
SUBTLE = RGBColor(0x2E, 0x5C, 0x8A)   # lighter blue, sub-headings
INK = RGBColor(0x20, 0x20, 0x20)

OUT = Path(r"e:\Claude\Dissertation\financial-proposal-platform\MidSem_Report_Garvit_2024AA05088.docx")

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


# ---------------------------------------------------------------- helpers ----
def _field(paragraph, instr):
    """Insert a Word field (used for the page-number footer)."""
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve"); instr_el.text = instr
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr_el); run._r.append(e)


def heading(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. {text}" if num else text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    # thin rule under heading
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def sub(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = SUBTLE
    return p


def para(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(" — " + text)
    else:
        p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ------------------------------------------------------------- header/footer -
sec = doc.sections[0]
sec.header.is_linked_to_previous = False
hp = sec.header.paragraphs[0]
hp.text = ""
hr = hp.add_run("Mid-Semester Report | AIMLCZG628T | 2024AA05088")
hr.font.size = Pt(8); hr.font.color.rgb = ACCENT
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

fp = sec.footer.paragraphs[0]
fp.text = ""
fr = fp.add_run("BITS Pilani WILP | M.Tech. AIML Dissertation        Page ")
fr.font.size = Pt(8)
_field(fp, "PAGE")
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ===================================================================== COVER =
def cover_line(text, size, bold=True, after=6, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

doc.add_paragraph().paragraph_format.space_after = Pt(18)
cover_line("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI", 14, color=ACCENT)
cover_line("Work Integrated Learning Programmes (WILP) Division", 11, bold=False, after=2)
cover_line("Second Semester, Academic Year 2025–2026", 11, bold=False, after=20)
cover_line("MID-SEMESTER REPORT", 18, color=ACCENT, after=18)
cover_line("Agentic SLM-Based Multi-Source Intelligent Proposal Generation System", 13)
cover_line("for Enterprise Consulting Using Adaptive Model Orchestration", 13, after=22)

ident = [
    ("Course Number", "AIMLCZG628T"),
    ("Course Title", "Dissertation"),
    ("Student Name", "Garvit D Jain"),
    ("BITS ID", "2024AA05088"),
    ("Degree Programme", "M.Tech. (WILP) – Artificial Intelligence & Machine Learning"),
    ("Research Area", "Agentic AI Systems with a Mixture of Specialised Small Language "
                      "Models for Enterprise Natural Language Processing"),
    ("Organization", "Virtusa Consultancy Services Pvt. Ltd., Chennai"),
    ("Supervisor", "Srinivas Ganesh, Associate Architect, Virtusa Consultancy Services "
                   "Pvt. Ltd., Hyderabad"),
    ("Additional Examiner", "Srivatsan Thirumalai, Technical Java Project Manager, "
                            "Virtusa Consultancy Services Pvt. Ltd., Chennai"),
    ("Report Date", "June 2026"),
]
t = doc.add_table(rows=0, cols=2)
t.style = "Light List Accent 1"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in ident:
    cells = t.add_row().cells
    cells[0].text = ""
    rk = cells[0].paragraphs[0].add_run(k); rk.bold = True; rk.font.size = Pt(10.5)
    cells[1].text = ""
    rv = cells[1].paragraphs[0].add_run(v); rv.font.size = Pt(10.5)
    cells[0].width = Pt(150)
from docx.shared import Inches
for row in t.rows:
    row.cells[0].width = Inches(1.9)
    row.cells[1].width = Inches(4.4)

doc.add_paragraph().paragraph_format.space_after = Pt(30)
cover_line("June 2026", 11, bold=False)

doc.add_page_break()

# =================================================================== ABSTRACT =
heading(None, "Abstract")
para(
    "Drafting a client proposal in an enterprise consulting practice is still a "
    "largely manual exercise. A consultant gathers fragments from issue trackers, "
    "earlier engagement documents, customer-relationship records and analyst notes, "
    "then reconciles them into a single persuasive narrative. The work is slow, the "
    "house voice drifts between authors, and the provenance of any number that ends "
    "up in the document is rarely recorded. General-purpose large language models "
    "(LLMs) can accelerate the writing, but doing so through a hosted frontier model "
    "raises three problems for a consulting firm: confidential client material must "
    "be sent across an external boundary, the per-request cost is hard to predict at "
    "volume, and the model offers no contractual guarantee that a quoted figure was "
    "actually drawn from a trusted source rather than invented."
)
para(
    "This dissertation designs and implements an agentic proposal-generation system "
    "that addresses those problems by orchestrating a mixture of small language "
    "models (SLMs) behind a governed retrieval pipeline, and by running the entire "
    "workflow inside the enterprise network with no outbound connectivity. The "
    "system reads a free-text request, infers the structured brief from it, retrieves "
    "supporting material across three independently governed knowledge repositories "
    "(financial evidence, prior-proposal exemplars and document templates), ranks "
    "each repository separately, and only proceeds to write when the factual grounding "
    "clears an explicit threshold. Every figure produced is checked back against a "
    "cited source; a number that cannot be traced is blocked, and when no admissible "
    "evidence exists the system emits a deliberately figure-free draft rather than "
    "fabricating one. A complete, replayable lineage record is persisted for each run "
    "and surfaced through a metrics dashboard and a prompt-history view."
)
para(
    "A working proof of concept has been built end to end. It is served fully offline "
    "by a local Qwen2.5-3B-Instruct model through an Ollama runtime, with a gateway "
    "abstraction that allows the same pipeline to be re-pointed at the firm's internal "
    "production SLM without code changes. This mid-semester report documents the "
    "completed design and development work — the multi-repository architecture, the "
    "adaptive model-orchestration seam, the anti-fabrication guardrails, the "
    "attachment-driven content path and the analytics surface — and sets out the "
    "evaluation and consolidation activities that remain for the testing phase."
)

doc.add_page_break()

# ==================================================================== CONTENTS
heading(None, "Contents")
contents = [
    "Broad Area of Work",
    "Background & Motivation",
    "Objectives",
    "Literature Review Summary",
    "System Data Sources & Knowledge Base",
    "Methodology & System Architecture Design",
    "Model Flow, Parameters, Memory & Compute",
    "Tools & Technology Stack",
    "Metrics & Prompt-History Subsystem",
    "Comparative Advantage over General-Purpose LLMs",
    "Work Completed to Date",
    "Plan of Work — Updated Status",
    "Key Findings & Observations",
    "Plan for the Remaining Semester",
    "References",
]
for i, c in enumerate(contents, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.add_run(f"{i}.  ").bold = True
    p.add_run(c)

doc.add_page_break()

# ============================================================ 1 BROAD AREA ====
heading(1, "Broad Area of Work")
para(
    "This dissertation lies where three currents in applied artificial intelligence "
    "meet: autonomous (agentic) AI workflows, the rise of compact task-specialised "
    "language models, and retrieval-grounded text generation for regulated business "
    "settings. The unifying problem is the automated authoring of enterprise "
    "consulting proposals from heterogeneous internal sources, performed in a way "
    "that an organisation can trust, afford and audit. The specific strands the work "
    "touches are:"
)
bullet("the design of cooperating software agents that each own a narrow step of the "
       "drafting workflow — understanding the request, retrieving knowledge, writing, "
       "and checking the result.", "Agentic AI orchestration")
bullet("the use of small, efficient language models in place of a single large model, "
       "so that inference can run on modest internal hardware at predictable cost.",
       "Small Language Models (SLMs)")
bullet("grounding generated text in an enterprise's own governed corpora so that "
       "claims are evidenced rather than improvised.",
       "Retrieval-Augmented Generation (RAG)")
bullet("a routing seam that selects the appropriate model for the environment and the "
       "task, and that can be re-pointed from a development model to the firm's "
       "production model without rewriting the pipeline.",
       "Adaptive model orchestration")
bullet("verification and provenance machinery that traces every quantitative claim to "
       "a cited source and records a replayable history of each generation.",
       "Governance, auditability & tool integration")

# ============================================================ 2 BACKGROUND ====
heading(2, "Background & Motivation")
para(
    "Virtusa, the employing organisation for this dissertation, delivers consulting "
    "and technology engagements across many client accounts. Each new opportunity "
    "typically begins with a written proposal that must read in the firm's voice, "
    "reflect the client's situation accurately, and respect commercial confidentiality. "
    "The raw material for such a document already exists inside the organisation — in "
    "Jira histories, in the archive of proposals that have previously won or lost, in "
    "CRM records, and in approved document templates — yet assembling it remains a "
    "repetitive task that consumes senior time and produces uneven results."
)
para(
    "Two industry shifts make automating this task newly feasible. First, language "
    "models have become competent enough at multi-step reasoning to be combined into "
    "agentic systems that plan and act rather than merely answer a single prompt. "
    "Second, a generation of small language models has emerged that approaches the "
    "quality of much larger models on narrow, well-scoped tasks while costing a "
    "fraction as much to run. Together these trends suggest an alternative to calling "
    "one expensive general model for everything: a coordinated set of smaller, "
    "specialised models, each applied where it is strongest."
)
para(
    "A naive deployment of a hosted frontier model into this setting would, however, "
    "be unacceptable for a consulting firm handling material non-public information. "
    "Client financials cannot be transmitted to an external service; spend cannot "
    "scale linearly with the number of proposals drafted; and a model that "
    "confidently states an unverifiable figure is worse than no model at all, because "
    "a plausible-looking but unsourced number can survive review and reach a client. "
    "The motivation for this project is therefore not only to draft proposals "
    "automatically, but to do so under hard constraints of confidentiality, cost and "
    "verifiability — constraints that shape every architectural decision described "
    "later in this report."
)

# ============================================================ 3 OBJECTIVES ====
heading(3, "Objectives")
para("The dissertation pursues the following objectives:")
for o in [
    "Survey the current state of agentic AI systems, small language models and "
    "retrieval-augmented generation as they apply to enterprise document automation.",
    "Design a multi-agent architecture in which distinct agents handle requirement "
    "understanding, knowledge retrieval, drafting and verification for proposal "
    "generation.",
    "Build a system that combines specialised small language models with a fallback "
    "to larger models where genuinely needed, to produce high-quality consulting "
    "proposals.",
    "Integrate multiple enterprise sources — issue trackers, historical proposals, "
    "templates and external knowledge — as governed, separately ranked inputs.",
    "Develop an adaptive model-orchestration mechanism (a model router) that selects "
    "the appropriate model for a task and deployment environment under cost and "
    "performance considerations.",
    "Enable tool calling and external integration where the deployment policy permits "
    "it, while preserving the air-gapped default.",
    "Define an evaluation scheme that assesses a generated proposal on relevance, "
    "completeness and consistency, backed by per-run quality metrics.",
    "Provide a feedback-aware path for improving outputs over time through editable, "
    "versioned drafts.",
    "Deliver a usable interface and a working proof of concept of the complete system.",
]:
    bullet(o)

# ============================================================ 4 LIT REVIEW ====
heading(4, "Literature Review Summary")
sub("4.1  Small Language Models for Agentic Systems")
para(
    "A recurring argument in recent work is that agentic workloads do not require a "
    "single very large model. Belcak et al. [1] contend that small language models "
    "are the natural substrate for agentic AI, because agent steps are typically "
    "narrow, repetitive and latency-sensitive — exactly the regime where a compact "
    "model specialised for the step is more economical than a general one. Sharma et "
    "al. [7] examine the architectural trade-offs of building agent stacks from small "
    "models, and Nardien et al. [2] show that the capabilities of a large agent can "
    "be distilled into smaller models when retrieval and tool use are added back, "
    "recovering much of the lost competence at far lower serving cost."
)
sub("4.2  Multi-Agent Collaboration and Continual Improvement")
para(
    "Several frameworks address how multiple smaller models can cooperate. Liang et "
    "al. [3] propose CMAT, a collaboration-tuning approach that strengthens small "
    "language models by having them work together rather than in isolation. Saeidi et "
    "al. [4] introduce a failure-aware meta-agentic framework that treats the failure "
    "of an individual agent as a signal to re-route or retry, which is directly "
    "relevant to the grounding-loop and block-and-regenerate behaviour used in this "
    "project. Zhuang et al. [6] show that an agent's underlying capabilities can be "
    "lifted through continual pre-training, pointing toward a route for specialising "
    "the proposal-writing model on a firm's own corpus."
)
sub("4.3  Retrieval-Augmented and Agentic RAG")
para(
    "Retrieval-augmented generation grounds a model's output in retrieved documents "
    "rather than in parametric memory alone; the foundational formulation is due to "
    "Lewis et al. [11]. Mishra et al. [5] survey the more recent agentic variants of "
    "RAG, in which retrieval is itself an action an agent decides to take, classifying "
    "the architectures and the ways they are evaluated. The present work adopts a "
    "deliberately strict form of grounding: retrieval is federated across three "
    "separately governed repositories, and only one of them is ever permitted to "
    "supply a citable figure."
)
sub("4.4  Foundation Models and Language Understanding")
para(
    "The broader lineage of the field underpins these choices. The transformer "
    "architecture of Vaswani et al. [12] is the common basis for the models used "
    "here; Devlin et al. [9] established bidirectional pre-training for language "
    "understanding; Brown et al. [10] demonstrated that scale yields few-shot "
    "competence; and Touvron et al. [8] showed that open, efficient foundation models "
    "can be both capable and practical to self-host — the property this project "
    "relies on to keep generation inside the enterprise boundary."
)

# ============================================================ 5 DATA SOURCES ==
heading(5, "System Data Sources & Knowledge Base")
para(
    "Rather than a single training dataset, this system is organised around three "
    "knowledge repositories, each with a distinct role and its own governance rules. "
    "Keeping them separate — and never merging their results into one ranking — is a "
    "core design invariant, because it is what lets the system reason differently "
    "about facts, style and structure."
)
table(
    ["Repository", "Role", "May supply…", "Governance"],
    [
        ["Financial (repo_financial)", "Evidence",
         "Citable figures and facts",
         "The only citable source; entity/period metadata must match the brief"],
        ["Proposals (repo_proposals)", "Exemplar",
         "House voice and phrasing only",
         "Anonymised on intake; never contributes a figure"],
        ["Templates (repo_templates)", "Scaffold",
         "Section structure and placeholders",
         "Structure only; verbatim placeholders preserved"],
    ],
)
para(
    "Documents enter these repositories through an ingestion pipeline that extracts "
    "text and tables, normalises and redacts personal and market-sensitive "
    "information, classifies the document, runs a quality gate, attaches metadata, "
    "chunks the content (keeping tables atomic), embeds each chunk and indexes it into "
    "the correct repository. Proposal exemplars pass an additional anonymisation gate "
    "that blocks any residual hard figure, grouped number, personal identifier or "
    "known client name before indexing — enforcing the rule that an exemplar teaches "
    "the system how to phrase a proposal, never what numbers to put in it. For "
    "development and demonstration the repositories are seeded in-process so that "
    "retrieval grounds without external services; the production deployment points the "
    "same code at a persistent vector store and a real embedding service."
)

# ============================================================ 6 ARCHITECTURE ==
heading(6, "Methodology & System Architecture Design")
sub("6.1  Clean-architecture layering")
para(
    "The backend follows a clean, dependency-inverted layering. A pure domain layer "
    "holds the entities and the port interfaces and performs no input or output. "
    "Around it sit use-case modules, one per bounded context (ingestion, retrieval, "
    "generation, reporting, metrics). Infrastructure adapters implement the ports — "
    "the vector store, the object store, the embedder and the language-model gateway "
    "— and an API layer exposes the use cases over HTTP. Dependencies point inward "
    "only, so the same generation logic runs unchanged whether it is wired to "
    "in-process test doubles or to live production services."
)
sub("6.2  The agentic generation pipeline")
para(
    "A generation request flows through a sequence of cooperating stages, each of "
    "which behaves like a specialised agent with a single responsibility:"
)
for s in [
    ("Intake / brief extraction", "a model-backed extractor reads the free-text "
     "request and infers the structured brief — the entity, fiscal year, proposal "
     "type, sector and requested line items. A fast regular-expression path resolves "
     "plainly named companies without a model call; the model is consulted only when "
     "that fails. Inferred values fill gaps only and never override an explicit "
     "choice, and an inferred value is treated as a guess, not as evidence."),
    ("Query formulation", "the brief is rewritten into one tailored query per "
     "repository, with the financial query leading on entity, period and line items."),
    ("Federated retrieval (fan-out)", "the three repositories are searched "
     "concurrently, each returning a role-tagged candidate pool."),
    ("Within-repository ranking", "each repository is ranked on its own terms; "
     "financial candidates whose entity or period do not match are dropped, won "
     "proposals are boosted as exemplars, and the best-matching template is selected."),
    ("Grounding gate (loop or refuse)", "the strength of the financial grounding is "
     "measured; if it falls below the floor the system widens financial recall and "
     "retries up to a bounded number of times, then either falls back to a figure-free "
     "draft or refuses outright."),
    ("Context assembly", "the surviving material is packed into the model's context "
     "window under a fixed budget that guarantees evidence the largest share, with "
     "bounded shares for exemplars and the scaffold."),
    ("Generation", "the language model writes the proposal section by section, "
     "instructed to mirror the exemplar voice and to state no figure that is not "
     "present in the supplied evidence."),
    ("Numeric verification + factual-health guard", "every figure in the draft is "
     "traced back to a cited financial chunk; an untraceable figure blocks the draft, "
     "and the share of factual content that resolves to the financial repository must "
     "be essentially total."),
    ("Confidence, contribution & lineage", "a confidence band is computed, the "
     "contribution of each repository is calculated, and a full generation event is "
     "persisted — always, even on a refusal — so the run can be replayed and audited."),
]:
    bullet(s[1], s[0])
para(
    "Three terminal behaviours fall out of this design. When grounding is strong the "
    "system produces a cited proposal. When there is no admissible evidence but a "
    "template and exemplars exist, it produces a figure-free \"style-only\" draft that "
    "uses bracketed placeholders where a number would go, so a consultant receives a "
    "correctly structured starting point instead of a dead end. When a client file is "
    "attached, a dedicated path turns that file into a proposal with a section count "
    "driven by the material itself and reproduces the client's own data tables "
    "verbatim under their original captions — allowing the figures the client "
    "supplied while still blocking any figure invented from neither the attachment "
    "nor cited evidence."
)
sub("6.3  Adaptive model orchestration")
para(
    "The dissertation's \"adaptive model orchestration\" is realised at the language-"
    "model gateway seam. Every component that needs text talks to a single gateway "
    "port; a factory chooses the concrete provider from configuration. In the local "
    "environment it returns either a deterministic echo stub (for tests), the local "
    "Ollama model (for realistic offline prose) or, where explicitly enabled, an "
    "external model for comparison. Outside development it returns the firm's internal "
    "SLM, and the settings layer fails closed so that a production deployment can "
    "never select an external provider. Because the routing decision is centralised "
    "and configuration-driven, swapping the development model for the production SLM — "
    "or introducing additional task-specialised models behind the same port — is a "
    "configuration change, not a rewrite. The intake extractor and the drafting stage "
    "already exercise this seam as separate calls, which is the foundation on which "
    "per-task model selection is being extended during the testing phase."
)

# ====================================================== 7 PARAMS/MEMORY/COMPUTE
heading(7, "Model Flow, Parameters, Memory & Compute")
para(
    "The proof of concept runs entirely on commodity, CPU-only hardware with no GPU "
    "acceleration and no network egress. The serving model is Qwen2.5-3B-Instruct, "
    "hosted by a local Ollama runtime; a smaller Qwen2.5-1.5B variant is also "
    "available as a faster fallback. The table below records the operative parameters "
    "as configured in the platform."
)
table(
    ["Parameter", "Value", "Why"],
    [
        ["Serving model", "Qwen2.5-3B-Instruct (~1.9 GB)",
         "Best quality that fits CPU-only inference on the dev box"],
        ["Fallback model", "Qwen2.5-1.5B (~0.99 GB)", "Lower latency when needed"],
        ["Runtime", "Ollama on localhost:11434", "Fully offline, no API key, no egress"],
        ["Context window (num_ctx)", "8192 tokens",
         "Matches the assembler's budget; avoids silent truncation"],
        ["Max output / section (num_predict)", "512 tokens",
         "Caps CPU generation time (pipeline requests 1024)"],
        ["Temperature", "0.0", "Deterministic output; discourages fabrication"],
        ["Keep-alive", "up to 24h", "Keeps the model resident; avoids cold reloads"],
        ["Request timeout", "300 s", "Tolerates CPU-bound multi-section runs"],
        ["Embedding", "256-dim local hashing embedder",
         "Deterministic, offline; pinned version for corpus comparability"],
        ["Hardware", "CPU-only, 16 GB RAM",
         "Demonstrates feasibility on modest internal compute"],
    ],
)
sub("7.1  Retrieval & grounding parameters")
para(
    "The governance knobs that shape every run are centralised in policy objects and "
    "recorded with each generation event:"
)
bullet("financial 40, proposal 8, template 1 candidates per fan-out branch.",
       "Branch recall (k)")
bullet("8192 tokens, split 60% evidence / 25% exemplars / 15% scaffold.",
       "Context budget")
bullet("a minimum financial grounding strength of 0.60 to proceed, 0.80 for the "
       "high-confidence band, up to 2 grounding loops before fallback or refusal.",
       "Grounding")
bullet("at least 99.9% of factual content must resolve to the financial repository, "
       "or the draft is blocked and regenerated.", "Factual-health floor")
sub("7.2  Observed performance")
para(
    "With the model pre-warmed and resident, a three-section advisory proposal "
    "completes in roughly 190–210 seconds on CPU; a single-section figure-free draft "
    "completes in about a minute; and the attachment-driven path completes in a "
    "comparable 190–210 seconds while emitting a dynamic number of sections. These "
    "figures confirm that grounded, governed generation is practical on internal "
    "hardware without specialised accelerators, which is the central feasibility claim "
    "of the dissertation."
)

# ============================================================ 8 TOOLS =========
heading(8, "Tools & Technology Stack")
table(
    ["Category", "Tools / Frameworks"],
    [
        ["Language model serving", "Ollama · Qwen2.5-3B / 1.5B-Instruct (local)"],
        ["Backend", "Python 3.12 · FastAPI · async SQLAlchemy 2.0 · Pydantic"],
        ["Orchestration", "Plain asyncio fan-out/fan-in pipeline (LangGraph-shaped)"],
        ["Retrieval & vectors", "ChromaDB (3 collections) · local hashing embedder"],
        ["Persistence", "PostgreSQL (catalogue, lineage, audit) · Alembic migrations"],
        ["Document I/O", "PyMuPDF · python-docx · xhtml2pdf / ReportLab (export)"],
        ["Frontend", "Next.js 15 · React 18 · TypeScript 5 · Tailwind CSS"],
        ["Deployment", "Docker (multi-stage) · Kubernetes (kustomize) · NetworkPolicies"],
        ["Testing", "pytest (62 backend tests) · live Postgres E2E · manifest lint"],
    ],
)

# ============================================================ 9 METRICS =======
heading(9, "Metrics & Prompt-History Subsystem")
para(
    "Because trust is a first-class requirement, the system treats observability as "
    "part of the product rather than an afterthought. Two operator-facing views expose "
    "what the system retrieved, wrote and decided."
)
sub("9.1  Metrics dashboard")
para(
    "The metrics page is organised into three zones. The first reports the composition "
    "of the knowledge base: counts of financial documents, proposal examples and "
    "templates, the total number of embedded chunks, the timestamp of the last "
    "ingestion, and the percentage each repository contributes to the corpus. The "
    "second zone summarises generation health over a rolling window (seven days by "
    "default): average confidence, average extraction quality, the refusal rate and "
    "the number of proposals produced, together with a runs-per-day bar chart and a "
    "donut chart that buckets runs by information loss into low, medium and high "
    "bands. The third zone embeds a table of the most recent runs, each of which links "
    "through to its full execution report."
)
sub("9.2  Prompt-history page")
para(
    "The prompt-history page lists every generation run, newest first, with a free-"
    "text search and status filter. Each row records the originating prompt and its "
    "proposal, the timestamp, the number of distinct files used, the outcome "
    "(generated, style-only or refused), the total processing time, the OCR and "
    "extraction quality of the documents retrieved, the information-loss percentage, "
    "and the repository contribution. Refused runs are still listed in full, with the "
    "quality columns shown as not-applicable — a deliberate choice, because a refusal "
    "is itself an auditable event."
)
sub("9.3  Execution report (drill-in)")
para(
    "Selecting any run opens its execution report, which reconstructs the run in ten "
    "sections: the verbatim prompt, the files used, the financial, proposal and "
    "template documents that were retrieved with their scores, the OCR confidence, the "
    "extraction quality, the information loss with the gate verdict, the stage-by-"
    "stage timing, and the final citations resolved to source and page. This is the "
    "evidence trail that distinguishes the system from an opaque chat interface: a "
    "reviewer can see not only what was written, but why, and on what basis."
)

# =================================================== 10 COMPARATIVE ADVANTAGE ==
heading(10, "Comparative Advantage over General-Purpose LLMs")
para(
    "The value of this system is clearest in scenarios where simply calling a hosted "
    "general-purpose model such as a frontier chat LLM would be unsuitable. Three "
    "representative use cases illustrate the difference."
)
sub("10.1  Drafting against confidential client financials")
para(
    "Consider a proposal that must quote a client's non-public revenue and margin. "
    "Sending that data to an external model would breach the firm's confidentiality "
    "obligations and potentially data-protection rules. This system never leaves the "
    "enterprise boundary: the model is served locally, retrieval runs against internal "
    "repositories, and the deployment fails closed if any configuration would enable "
    "egress. A general hosted LLM cannot offer that guarantee by construction."
)
sub("10.2  High-volume drafting at predictable cost")
para(
    "A practice that produces hundreds of proposals a month faces linear, "
    "hard-to-forecast spend if each draft is billed per token by an external model. "
    "Running compact models on existing CPU hardware turns that variable cost into a "
    "largely fixed one, and the mixture-of-small-models design means a routine step is "
    "handled by a cheap specialised model rather than an expensive general one. The "
    "demonstrated CPU-only latency shows this is operationally realistic rather than "
    "merely cheaper on paper."
)
sub("10.3  Verifiable, refuse-rather-than-fabricate output")
para(
    "A general LLM asked for a figure it does not know will often supply a "
    "confident-sounding but unsourced number, with no provenance and no option to "
    "decline. This system inverts that default: a figure that cannot be traced to a "
    "cited financial source is blocked, an ungrounded request yields a figure-free "
    "draft or an explicit refusal, and every decision is recorded in a replayable "
    "lineage. For a document that will be put in front of a client, a defensible "
    "\"I cannot evidence this\" is far more valuable than a fluent guess."
)
para(
    "A fourth, longer-term advantage follows from the architecture: because the "
    "drafting model sits behind a single routing seam, the firm can specialise and "
    "update small models on its own corpus far more cheaply than it could fine-tune or "
    "replace a monolithic model — the direction set out in the remaining-work plan."
)

# ============================================================ 11 WORK DONE =====
heading(11, "Work Completed to Date")
table(
    ["Activity", "Completion %", "Remarks"],
    [
        ["Literature survey & problem scoping", "100%", "Completed"],
        ["Architecture & domain design (clean architecture, ports)", "100%", "Completed"],
        ["Three-repository ingestion + governance gates", "100%", "Completed & tested"],
        ["RAG retrieval, grounding & generation pipeline", "100%", "Completed & tested"],
        ["Anti-fabrication guardrails (numeric + factual health)", "100%", "Completed & tested"],
        ["Adaptive model-orchestration seam (gateway/factory)", "100%", "Completed"],
        ["Local SLM serving (Ollama, offline)", "100%", "Completed"],
        ["Attachment-driven content path", "100%", "Completed & verified"],
        ["API, web UI, metrics & history surface", "100%", "Completed"],
        ["Export (PDF/DOCX, company template)", "100%", "Completed"],
        ["Deployment hardening (Docker, K8s, air-gap policies)", "100%", "Completed"],
        ["Quantitative evaluation & user study", "20%", "In progress (testing phase)"],
        ["Per-task multi-SLM routing & specialisation", "30%", "In progress"],
        ["Dissertation write-up", "25%", "In progress"],
    ],
)
para(
    "In summary, the full proof of concept has been designed, implemented and brought "
    "up end to end: it ingests documents into three governed repositories, retrieves "
    "and grounds against them, generates governed proposals with a local offline "
    "model, blocks unverifiable figures, exports a company-styled document, and "
    "exposes a complete metrics and audit surface. The backend is covered by an "
    "automated test suite, and the deployment has been validated against a real "
    "PostgreSQL instance and a set of linted Kubernetes manifests. The emphasis now "
    "moves from building to measuring."
)

# ============================================================ 12 PLAN STATUS ===
heading(12, "Plan of Work — Updated Status")
table(
    ["Phase", "Timeline", "Status"],
    [
        ["Dissertation Outline", "25 Apr 2026 – 05 May 2026", "COMPLETED"],
        ["Design & Development", "06 May 2026 – 12 Jun 2026", "COMPLETED"],
        ["Testing & User Evaluation", "13 Jun 2026 – 18 Jul 2026", "IN PROGRESS"],
        ["Dissertation Review", "19 Jul 2026 – 26 Jul 2026", "PENDING"],
        ["Submission", "26 Jul 2026 – 30 Jul 2026", "PENDING"],
    ],
)
para(
    "As of this report (June 2026) the design and development phase is complete and "
    "the project has entered the testing and user-evaluation phase on schedule."
)

# ============================================================ 13 FINDINGS ======
heading(13, "Key Findings & Observations")
for f in [
    "Grounded, governed proposal generation is feasible entirely offline on CPU-only "
    "hardware. A 3B-parameter model produces usable advisory prose in roughly three "
    "minutes per multi-section draft, removing the need for a GPU or an external API.",
    "Keeping the three repositories separate is what makes the governance tractable. "
    "Because only the financial repository can be cited, the rules for facts, style "
    "and structure can be enforced independently and the provenance of every figure "
    "stays unambiguous.",
    "A strict numeric verifier paired with a figure-free fallback resolves the central "
    "tension of the project: the system can still be helpful when it has no evidence, "
    "by writing a correctly structured placeholder draft, without ever shipping an "
    "invented number.",
    "Centralising model selection behind one gateway seam makes the move from a "
    "development model to the production internal SLM — and the future addition of "
    "task-specialised models — a configuration change rather than a re-engineering "
    "effort.",
    "Treating the audit trail as a product surface, through the metrics dashboard and "
    "execution reports, is what converts an opaque generator into a tool a regulated "
    "firm can actually adopt.",
]:
    bullet(f)

# ============================================================ 14 REMAINING ====
heading(14, "Plan for the Remaining Semester")
sub("14.1  Testing & evaluation (until 18 July 2026)")
for x in [
    "Define and run a quantitative evaluation of generated proposals on relevance, "
    "completeness and consistency against a held-out set of briefs.",
    "Measure grounding accuracy, the fabricated-figure block rate and the refusal "
    "rate, and characterise latency across the grounded, style-only and "
    "attachment-driven paths.",
    "Conduct a small user evaluation with consulting practitioners to assess draft "
    "usefulness and editing effort.",
    "Extend the orchestration seam toward genuine per-task model routing, comparing a "
    "single general model against a mixture of specialised small models on cost and "
    "quality.",
]:
    bullet(x)
sub("14.2  Review & submission (19 July – 30 July 2026)")
for x in [
    "Author the full dissertation document and incorporate the evaluation results.",
    "Submit to the supervisor and additional examiner for review and feedback.",
    "Prepare the reproducible code artefact and finalise the report for submission.",
]:
    bullet(x)

# ============================================================ 15 REFERENCES ===
heading(15, "References")
refs = [
    "P. Belcak et al., \"Small Language Models are the Future of Agentic AI,\" arXiv "
    "preprint arXiv:2506.02153, 2025.",
    "Nardien et al., \"Distilling LLM Agents into Small Models with Retrieval and Tool "
    "Use,\" arXiv preprint arXiv:2505.17612, 2025.",
    "X. Liang et al., \"CMAT: A Multi-Agent Collaboration Tuning Framework for "
    "Enhancing Small Language Models,\" arXiv preprint arXiv:2404.01663, 2024.",
    "A. Saeidi et al., \"FAMA: Failure-Aware Meta-Agentic Framework for LLM-based "
    "Systems,\" arXiv preprint arXiv:2604.25135, 2026.",
    "S. Mishra et al., \"Agentic Retrieval-Augmented Generation (RAG): Taxonomy, "
    "Architectures, and Evaluation,\" arXiv preprint, 2026.",
    "Y. Zhuang et al., \"Hephaestus: Improving Agent Capabilities through Continual "
    "Pre-training,\" arXiv preprint arXiv:2502.06589, 2025.",
    "R. Sharma et al., \"Small Language Models for Agentic Systems: Architectures and "
    "Trade-offs,\" arXiv preprint, 2025.",
    "H. Touvron et al., \"LLaMA: Open and Efficient Foundation Language Models,\" arXiv "
    "preprint arXiv:2302.13971, 2023.",
    "J. Devlin et al., \"BERT: Pre-training of Deep Bidirectional Transformers for "
    "Language Understanding,\" NAACL, 2019.",
    "T. Brown et al., \"Language Models are Few-Shot Learners,\" NeurIPS, 2020.",
    "P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP "
    "Tasks,\" NeurIPS, 2020.",
    "A. Vaswani et al., \"Attention Is All You Need,\" NeurIPS, 2017.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"[{i}] ").bold = True
    p.add_run(r)

# signature block
doc.add_paragraph().paragraph_format.space_after = Pt(18)
sig = doc.add_table(rows=2, cols=3)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
names = ["Garvit D Jain", "Srinivas Ganesh", "Srivatsan Thirumalai"]
roles = ["Student | 2024AA05088", "Supervisor", "Additional Examiner"]
for i in range(3):
    c0 = sig.rows[0].cells[i]; c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = c0.paragraphs[0].add_run(names[i]); rn.bold = True; rn.font.size = Pt(10.5)
    c1 = sig.rows[1].cells[i]; c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = c1.paragraphs[0].add_run(roles[i]); rr.font.size = Pt(9.5)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"WROTE {OUT}  ({OUT.stat().st_size} bytes)")
