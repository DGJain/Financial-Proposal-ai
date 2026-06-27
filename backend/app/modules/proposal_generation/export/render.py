"""Server-side proposal renderers (ui-design.md §6.4).

Export renders the proposal as a **finished company-template document** — a clean
proposal in the firm's house style (header rule, cover title, numbered sections,
source endnotes), the way a consultant would send it out. It deliberately does
*not* stamp the run's metrics/lineage into the body; that provenance lives in the
Execution Report, not in the client-facing document.

Two flows share one canonical stylesheet:

* **From the stored proposal** — ``render_document_fragment`` lays the locked
  sections into the template; ``render_document_page`` wraps that with the CSS for
  the in-browser editor (``GET /proposals/{id}/document``).
* **From edited HTML** — the WYSIWYG editor posts back the HTML the user actually
  edited; ``render_pdf_from_html`` / ``render_docx_from_html`` convert *that* so the
  exported PDF/DOCX is exactly what was on screen.

Air-gap friendly: PDF via xhtml2pdf (pure-Python ReportLab backend), DOCX via a
small stdlib ``html.parser`` walk into python-docx. No network, no system libs.
"""

from __future__ import annotations

import os
from html import escape
from html.parser import HTMLParser

from app.domain.generation.generation_event import GenerationEvent
from app.domain.proposals.proposal import Proposal

# The firm whose house style the document wears. Placeholder until branded —
# override with EXPORT_FIRM_NAME (settings.py is linter-managed, so read env here).
FIRM_NAME = os.getenv("EXPORT_FIRM_NAME", "Halstead Partners")

# Canonical house stylesheet. Kept to the CSS subset xhtml2pdf understands (block
# + table layout, no flexbox/grid) so the on-screen editor and the exported PDF
# render identically. Mirrors the anonymised exemplar: a running firm header rule,
# a dark-red accent for section numbers, banded section spacing.
DOCUMENT_CSS = """
@page { size: A4; margin: 2.2cm 2cm; }
body { font-family: Georgia, "Times New Roman", serif; color: #1a1a1a; line-height: 1.5; font-size: 11pt; }
.doc { max-width: 52rem; margin: 0 auto; }
.doc .firm-rule { border-bottom: 2px solid #1f2a44; padding-bottom: 6px; margin-bottom: 24px;
  font-family: Arial, Helvetica, sans-serif; font-size: 9pt; letter-spacing: .08em; text-transform: uppercase; color: #1f2a44; }
.doc .firm-rule .name { font-weight: 700; }
.doc .firm-rule .kicker { color: #7a7f8a; }
.doc .cover { margin-bottom: 28px; }
.doc .eyebrow { font-family: Arial, Helvetica, sans-serif; font-size: 9pt; font-weight: 700;
  letter-spacing: .12em; text-transform: uppercase; color: #7a1f2b; margin-bottom: 6px; }
.doc h1.title { font-family: Arial, Helvetica, sans-serif; font-size: 26pt; line-height: 1.1;
  color: #1f2a44; margin: 0 0 8px 0; }
.doc .prepared { font-family: Arial, Helvetica, sans-serif; font-size: 11pt; color: #5b6170; }
.doc section.sec { margin-bottom: 18px; }
.doc h2 { font-family: Arial, Helvetica, sans-serif; font-size: 14pt; color: #1f2a44; margin: 18px 0 6px 0; }
.doc h2 .num { color: #7a1f2b; font-weight: 700; }
.doc .body p { margin: 0 0 8px 0; }
.doc table { border-collapse: collapse; width: 100%; margin: 8px 0; font-size: 10pt; }
.doc th { background: #1f2a44; color: #fff; text-align: left; padding: 6px 8px;
  font-family: Arial, Helvetica, sans-serif; font-size: 9pt; }
.doc td { border-bottom: 1px solid #e2e4ea; padding: 6px 8px; }
.doc .sources { margin-top: 28px; padding-top: 10px; border-top: 1px solid #ccc;
  font-family: Arial, Helvetica, sans-serif; font-size: 8.5pt; color: #5b6170; }
.doc .sources .h { font-weight: 700; text-transform: uppercase; letter-spacing: .08em; margin-bottom: 4px; }
.doc .doc-footer { margin-top: 22px; font-family: Arial, Helvetica, sans-serif; font-size: 8pt; color: #9aa0ad; }
""".strip()


def _looks_like_table(block: str) -> bool:
    """A blank-line-delimited block of ≥2 lines that each contain a pipe = a data table."""
    lines = [ln for ln in block.splitlines() if ln.strip()]
    return len(lines) >= 2 and all("|" in ln for ln in lines)


def _render_table(block: str) -> str:
    """Render a pipe-delimited block as an HTML table (first row = header)."""
    rows: list[list[str]] = []
    for ln in block.splitlines():
        if not ln.strip():
            continue
        cells = [c.strip() for c in ln.split("|")]
        # Drop empty edge cells from "| a | b |"-style rows.
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        # Skip Markdown separator rows ("---|:--:|---").
        if cells and all(set(c) <= set("-: ") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return f"<p>{escape(block)}</p>"
    head, *rest = rows
    thead = "".join(f"<th>{escape(c)}</th>" for c in head)
    body = "".join(
        "<tr>" + "".join(f"<td>{escape(c)}</td>" for c in r) + "</tr>" for r in rest
    )
    return f"<table><tr>{thead}</tr>{body}</table>"


def _paras(body: str) -> str:
    """Render a section body: blank-line-separated blocks become <p> (with <br/> for
    internal newlines), except pipe-delimited blocks, which become data <table>s."""
    blocks = [b.strip() for b in body.replace("\r\n", "\n").split("\n\n") if b.strip()]
    if not blocks:
        return "<p></p>"
    out: list[str] = []
    for b in blocks:
        if _looks_like_table(b):
            out.append(_render_table(b))
        else:
            out.append(f"<p>{escape(b).replace(chr(10), '<br/>')}</p>")
    return "".join(out)


def render_document_fragment(
    proposal: Proposal,
    event: GenerationEvent | None = None,
    *,
    firm_name: str = FIRM_NAME,
    title: str = "Proposal",
) -> str:
    """The proposal as a styled ``<div class="doc">`` — the editable/exportable body.

    No surrounding ``<style>``; pair with :func:`render_document_page` (editor) or the
    export wrapper (PDF/DOCX), both of which supply :data:`DOCUMENT_CSS`.
    """
    version = proposal.current_version
    firm = escape(firm_name)
    ordered = sorted(version.sections, key=lambda s: s.order)
    # A single-section document is a cover letter — it carries its own letterhead
    # (date, addressee, salutation, sign-off), so we render only the firm rule and the
    # flowing text, with no separate cover block and no numbered section headings.
    single = len(ordered) == 1
    parts: list[str] = [
        '<div class="doc">',
        f'<div class="firm-rule"><span class="name">{firm}</span>'
        f'<span class="kicker"> &nbsp;·&nbsp; Proposal</span></div>',
    ]
    if not single:
        parts += [
            '<div class="cover">',
            '<div class="eyebrow">Proposal</div>',
            f'<h1 class="title">{escape(title)}</h1>',
            f'<div class="prepared">Prepared by {firm}</div>',
            "</div>",
        ]
    for n, section in enumerate(ordered, start=1):
        parts.append('<section class="sec">')
        if not single:
            parts.append(f'<h2><span class="num">{n}.</span> {escape(section.heading)}</h2>')
        parts.append(f'<div class="body">{_paras(section.body)}</div>')
        parts.append("</section>")

    if event is not None and event.citations:
        parts.append('<div class="sources"><div class="h">Sources</div>')
        for c in event.citations:
            parts.append(f"<div>[{c.claim_ordinal}] {escape(c.source_name)}, p.{c.page}</div>")
        parts.append("</div>")

    parts.append(
        f'<div class="doc-footer">{firm} — proposal {escape(proposal.proposal_id)} '
        f"· version {version.version_no}</div>"
    )
    parts.append("</div>")
    return "\n".join(parts)


def _wrap_page(fragment: str, *, title: str) -> str:
    """Wrap a ``.doc`` fragment into a full HTML document carrying the house CSS."""
    return (
        "<!DOCTYPE html>\n"
        '<html lang="en"><head><meta charset="utf-8">'
        f"<title>{escape(title)}</title>"
        f"<style>{DOCUMENT_CSS}</style></head><body>\n{fragment}\n</body></html>"
    )


def render_html_document(html: str, *, title: str = "Proposal") -> str:
    """Wrap an edited ``.doc`` fragment into a full styled HTML page for export.

    Pass-through if ``html`` is already a full document.
    """
    return html if "<html" in html.lower() else _wrap_page(html, title=title)


def render_document_page(
    proposal: Proposal,
    event: GenerationEvent | None = None,
    *,
    firm_name: str = FIRM_NAME,
    title: str = "Proposal",
) -> str:
    """Full styled HTML page for the in-browser editor (``GET .../document``)."""
    fragment = render_document_fragment(proposal, event, firm_name=firm_name, title=title)
    return _wrap_page(fragment, title=title)


# --- format renderers (stored proposal) --------------------------------------


def render_markdown(
    proposal: Proposal,
    event: GenerationEvent | None = None,
    *,
    firm_name: str = FIRM_NAME,
    title: str = "Proposal",
) -> str:
    """Plain-text Markdown of the proposal — sections in order + source endnotes."""
    version = proposal.current_version
    lines: list[str] = [f"# {title}", "", f"*Prepared by {firm_name}*", ""]
    ordered = sorted(version.sections, key=lambda s: s.order)
    single = len(ordered) == 1
    for n, section in enumerate(ordered, start=1):
        if not single:
            lines.append(f"## {n}. {section.heading}")
            lines.append("")
        lines.append(section.body)
        lines.append("")
    if event is not None and event.citations:
        lines.append("### Sources")
        lines.append("")
        for c in event.citations:
            lines.append(f"- [{c.claim_ordinal}] {c.source_name}, p.{c.page}")
        lines.append("")
    return "\n".join(lines)


def render_html(
    proposal: Proposal,
    event: GenerationEvent | None = None,
    *,
    firm_name: str = FIRM_NAME,
    title: str = "Proposal",
) -> str:
    """Self-contained HTML document of the proposal (house style)."""
    return render_document_page(proposal, event, firm_name=firm_name, title=title)


def render_pdf(
    proposal: Proposal,
    event: GenerationEvent | None = None,
    *,
    firm_name: str = FIRM_NAME,
    title: str = "Proposal",
) -> bytes:
    """PDF of the stored proposal in the house template."""
    fragment = render_document_fragment(proposal, event, firm_name=firm_name, title=title)
    return render_pdf_from_html(fragment, title=title)


def render_docx(
    proposal: Proposal,
    event: GenerationEvent | None = None,
    *,
    firm_name: str = FIRM_NAME,
    title: str = "Proposal",
) -> bytes:
    """DOCX of the stored proposal in the house template."""
    fragment = render_document_fragment(proposal, event, firm_name=firm_name, title=title)
    return render_docx_from_html(fragment, title=title)


# --- HTML → PDF / DOCX (used by both stored render and edited-HTML export) -----


def render_pdf_from_html(html: str, *, title: str = "Proposal") -> bytes:
    """Convert an HTML ``.doc`` fragment (or full page) to PDF via xhtml2pdf.

    Accepts either a bare fragment or a full page; a fragment is wrapped with the
    house CSS so an edited document exports exactly as it appeared in the editor.
    """
    from io import BytesIO

    from xhtml2pdf import pisa  # lazy

    page = html if "<html" in html.lower() else _wrap_page(html, title=title)
    buffer = BytesIO()
    status = pisa.CreatePDF(src=page, dest=buffer, encoding="utf-8")
    if status.err:
        raise RuntimeError(f"PDF rendering failed ({status.err} error(s))")
    return buffer.getvalue()


class _DocxHtmlBuilder(HTMLParser):
    """Walk our controlled proposal HTML into a python-docx document.

    Handles the tags the template and the WYSIWYG editor emit: headings, paragraphs,
    bold/italic runs, line breaks, lists and tables. ``<style>``/``<script>`` content
    is skipped so an exported full page doesn't leak its CSS into the body. Tables are
    buffered row-by-row and materialised at ``</table>`` once their dimensions are known.
    """

    _HEADING = {"h1": 0, "h2": 1, "h3": 2, "h4": 3}

    def __init__(self, doc) -> None:  # noqa: ANN001 - docx Document
        super().__init__(convert_charrefs=True)
        self._doc = doc
        self._para = None  # current python-docx paragraph
        self._bold = 0
        self._italic = 0
        self._skip = 0  # depth inside <style>/<script>
        # Table buffer: rows -> cells -> text fragments (None when not in a table).
        self._rows: list[list[str]] | None = None
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def _ensure_para(self):  # noqa: ANN202
        if self._para is None:
            self._para = self._doc.add_paragraph()
        return self._para

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag in ("style", "script"):
            self._skip += 1
        elif tag in self._HEADING:
            self._para = self._doc.add_heading("", level=self._HEADING[tag])
        elif tag == "p":
            self._para = self._doc.add_paragraph()
        elif tag == "li":
            self._para = self._doc.add_paragraph(style="List Bullet")
        elif tag in ("b", "strong"):
            self._bold += 1
        elif tag in ("i", "em"):
            self._italic += 1
        elif tag == "br" and self._para is not None:
            self._para.add_run().add_break()
        elif tag == "table":
            self._rows = []
        elif tag == "tr" and self._rows is not None:
            self._row = []
        elif tag in ("td", "th") and self._row is not None:
            self._cell = []

    def handle_endtag(self, tag: str) -> None:
        if tag in ("style", "script"):
            self._skip = max(0, self._skip - 1)
        elif tag in self._HEADING or tag in ("p", "li"):
            self._para = None
        elif tag in ("b", "strong"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("i", "em"):
            self._italic = max(0, self._italic - 1)
        elif tag in ("td", "th") and self._cell is not None and self._row is not None:
            self._row.append(" ".join(self._cell))
            self._cell = None
        elif tag == "tr" and self._row is not None and self._rows is not None:
            self._rows.append(self._row)
            self._row = None
        elif tag == "table" and self._rows is not None:
            self._flush_table(self._rows)
            self._rows = None

    def _flush_table(self, rows: list[list[str]]) -> None:
        rows = [r for r in rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self._doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r, cells in enumerate(rows):
            for c, text in enumerate(cells):
                table.rows[r].cells[c].text = text

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        if self._cell is not None:
            stripped = data.strip()
            if stripped:
                self._cell.append(stripped)
            return
        text = data.strip("\n")
        if not text.strip():
            return
        run = self._ensure_para().add_run(text)
        run.bold = self._bold > 0
        run.italic = self._italic > 0


def render_docx_from_html(html: str, *, title: str = "Proposal") -> bytes:
    """Convert an HTML ``.doc`` fragment (or full page) to a .docx (python-docx)."""
    from io import BytesIO

    from docx import Document as DocxDocument  # lazy

    doc = DocxDocument()
    doc.core_properties.title = title
    builder = _DocxHtmlBuilder(doc)
    builder.feed(html)
    builder.close()
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
